"""
学习通自测题目抓取工具
实现登录、课程选择、自测列表获取功能
"""

import argparse
import base64
import re
import time
from dataclasses import dataclass
from typing import Dict, List, Optional, Union
from urllib import parse
from urllib.parse import urlparse
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from Crypto.Cipher import AES
from Crypto.Util.Padding import pad


@dataclass
class Course:
    course_id: str
    class_id: str
    cpi: str
    course_name: str
    course_url: str


@dataclass
class Exam:
    exam_id: str
    exam_name: str
    exam_status: str
    exam_url: str
    course_id: str


@dataclass
class Question:
    answer_type: int
    question_title: str
    correct_answer: Union[str, List[str]]
    question_answers: Optional[List[str]] = None


class AESCrypto:
    def __init__(self, key: str = "u2oh6Vu^HWe4_AES"):
        self.key = key.encode('utf-8')

    def encrypt(self, message: str) -> str:
        cipher = AES.new(self.key, AES.MODE_CBC, self.key)
        padded_message = pad(message.encode('utf-8'), AES.block_size, style='pkcs7')
        encrypted_bytes = cipher.encrypt(padded_message)
        return base64.b64encode(encrypted_bytes).decode('utf-8')


class ExamCrawler:
    ANSWER_TYPES = {
        "单选题": 1,
        "多选题": 2,
        "填空题": 3,
        "判断题": 4,
        "名词解释": 6,
        "简答题": 7,
        "论述题": 7,
    }

    API_ENDPOINTS = {
        'login': 'http://passport2.chaoxing.com/fanyalogin',
        'course_list': 'https://mooc2-ans.chaoxing.com/mooc2-ans/visit/courselistdata',
        'course_middle': 'https://mooc1.chaoxing.com/visit/stucoursemiddle',
    }

    def __init__(self):
        self.session = requests.Session()
        self.crypto = AESCrypto()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
            '(KHTML, like Gecko) Chrome/132.0.0.0 Safari/537.36'
        })

    def login(self, phone: str, password: str) -> bool:
        try:
            encrypted_phone = self.crypto.encrypt(phone)
            encrypted_password = self.crypto.encrypt(password)
            headers = {
                "Accept": "application/json, text/javascript, */*; q=0.01",
                "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
                "Origin": "https://passport2.chaoxing.com",
                "Referer": "https://passport2.chaoxing.com/login",
                "X-Requested-With": "XMLHttpRequest"
            }
            payload = {
                'uname': encrypted_phone,
                'password': encrypted_password,
                't': "true",
                'validate': None,
                'doubleFactorLogin': "0",
                'independentId': "0"
            }
            response = self.session.post(
                self.API_ENDPOINTS['login'],
                data=payload,
                headers=headers,
                timeout=30
            )
            result = response.json()
            return result.get('status', False)
        except:
            return False

    def get_courses(self) -> List[Course]:
        try:
            params = {"courseType": "1", "courseFolderId": "0", "superstarClass": "0"}
            response = self.session.get(
                self.API_ENDPOINTS['course_list'],
                params=params,
                timeout=30
            )
            soup = BeautifulSoup(response.text, "lxml")
            courses = []
            for course_div in soup.find_all("div", attrs={"class": "course-info"}):
                link = course_div.find_next("a", attrs={"class": "color1"})
                if not link:
                    continue
                course_url = link.get("href", "")
                name_span = link.find_next("span", attrs={"class": "course-name overHidden2"})
                if not name_span:
                    continue
                course_name = name_span.get_text(strip=True)
                parsed_url = urlparse(course_url)
                url_data = parse.parse_qs(parsed_url.query)
                course = Course(
                    course_id=url_data.get("courseid", [""])[0],
                    class_id=url_data.get("clazzid", [""])[0],
                    cpi=url_data.get("cpi", [""])[0],
                    course_name=course_name,
                    course_url=course_url
                )
                if course.course_id and course.class_id:
                    courses.append(course)
            return courses
        except:
            return []

    def _get_examEnc(self, course: Course) -> str:
        try:
            params = {
                "courseid": course.course_id,
                "clazzid": course.class_id,
                "cpi": course.cpi,
                "ismooc2": 1,
                "v": time.time(),
            }
            response = self.session.get(
                self.API_ENDPOINTS['course_middle'],
                params=params,
                timeout=30
            )
            soup = BeautifulSoup(response.text, "lxml")
            input_tag = soup.find("input", id="examEnc")
            return str(input_tag.get("value", "")) if input_tag else ""
        except:
            return ""

    def _get_enc(self, course: Course) -> str:
        try:
            params = {
                "courseid": course.course_id,
                "clazzid": course.class_id,
                "cpi": course.cpi,
                "ismooc2": 1,
                "v": time.time(),
            }
            response = self.session.get(
                self.API_ENDPOINTS['course_middle'],
                params=params,
                timeout=30
            )
            soup = BeautifulSoup(response.text, "lxml")
            input_tag = soup.find("input", id="enc")
            return str(input_tag.get("value", "")) if input_tag else ""
        except:
            return ""

    def _get_openc(self, course: Course) -> str:
        try:
            params = {
                "courseid": course.course_id,
                "clazzid": course.class_id,
                "cpi": course.cpi,
                "ismooc2": 1,
                "v": time.time(),
            }
            response = self.session.get(
                self.API_ENDPOINTS['course_middle'],
                params=params,
                timeout=30
            )
            soup = BeautifulSoup(response.text, "lxml")
            input_tag = soup.find("input", id="openc")
            return str(input_tag.get("value", "")) if input_tag else ""
        except:
            return ""

    def get_exams(self, course: Course) -> List[Exam]:
        exams = []
        exam_urls_seen = set()
        exam_enc = self._get_examEnc(course)
        openc = self._get_openc(course)
        stuenc = self._get_enc(course)
        api_url = "https://mooc1.chaoxing.com/exam-ans/mooc2/exam/exam-list"
        params = {
            "courseid": course.course_id,
            "clazzid": course.class_id,
            "cpi": course.cpi,
            "ut": "s",
            "stuenc": stuenc,
            "enc": exam_enc,
            "openc": openc,
            "type": 1,
        }
        try:
            resp = self.session.get(api_url, params=params, timeout=30)
            soup = BeautifulSoup(resp.text, "lxml")
            for li in soup.find_all("li"):
                div = li.find("div", onclick=True)
                if div:
                    onclick = div.get("onclick", "")
                    match = re.search(r"viewPaper\(['\"](\d+)['\"]\)", onclick)
                    if match:
                        answer_id = match.group(1)
                        name_p = li.find("p", class_="overHidden2")
                        name = name_p.get_text(strip=True) if name_p else f"自测_{answer_id}"
                        exam_url = f"https://mooc1.chaoxing.com/exam-ans/exam/test/reVersionPaperMarkContentNew?courseId={course.course_id}&classId={course.class_id}&p=1&id={answer_id}&ut=s&cpi={course.cpi}&newMooc=true&openc={openc}"
                        if exam_url not in exam_urls_seen:
                            exam_urls_seen.add(exam_url)
                            exams.append(Exam(exam_id=answer_id, exam_name=name[:100], exam_status="", exam_url=exam_url, course_id=course.course_id))
        except:
            pass
        if not exams:
            print("未能自动获取自测列表，请粘贴自测页面URL：")
            manual_url = input().strip()
            if manual_url:
                parsed = urlparse(manual_url)
                qs = parse.parse_qs(parsed.query)
                answer_id = qs.get("id", [""])[0] or qs.get("examAnswerId", [""])[0]
                if answer_id:
                    exams.append(Exam(exam_id=answer_id, exam_name=f"自测_{answer_id}", exam_status="", exam_url=manual_url, course_id=course.course_id))
        return exams

    def get_exam_questions(self, exam: Exam) -> List[Question]:
        try:
            response = self.session.get(exam.exam_url, timeout=30)
            soup = BeautifulSoup(response.text, "lxml")
            mark_table = soup.find("div", class_="mark_table")
            if mark_table:
                question_divs = mark_table.find_all("div", class_="questionLi")
            else:
                question_divs = soup.find_all("div", class_="questionLi")
            if not question_divs:
                question_divs = soup.find_all("div", class_="singleQuesId")
            questions = []
            for q_div in question_divs:
                question = self._parse_question_div(q_div)
                if question:
                    questions.append(question)
            return questions
        except:
            return []

    def _parse_question_div(self, q_div) -> Optional[Question]:
        try:
            h3_tag = q_div.find("h3", class_="mark_name")
            if not h3_tag:
                h3_tag = q_div.find("h3")
                if not h3_tag:
                    return None
            h3_text = h3_tag.get_text(strip=True)
            answer_type = 1
            type_match = re.match(r'\d+\.\s*\((.*?)\)', h3_text)
            if type_match:
                q_type_str = type_match.group(1)
                for type_name, type_id in self.ANSWER_TYPES.items():
                    if type_name in q_type_str:
                        answer_type = type_id
                        break
            qt_content = h3_tag.find("span", class_="qtContent")
            if qt_content:
                ps = qt_content.find_all("p")
                if ps:
                    question_title = "\n".join(p.get_text(strip=True) for p in ps)
                else:
                    question_title = qt_content.get_text(strip=True)
            else:
                question_title = re.sub(r'^\d+\.\s*\(.*?\)\s*', '', h3_text).strip()
            if not question_title:
                question_title = h3_text
            options = []
            ul_tag = q_div.find("ul", class_="mark_letter")
            if ul_tag:
                for li in ul_tag.find_all("li"):
                    li_text = li.get_text(strip=True)
                    if li_text:
                        options.append(li_text)
            correct_answer = ""
            mark_answer = q_div.find("div", class_="mark_answer")
            if mark_answer:
                right_span = mark_answer.find("span", class_="rightAnswerContent")
                if right_span:
                    correct_answer = right_span.get_text(strip=True)
                else:
                    right_dd = mark_answer.find("dd", class_="rightAnswerContent")
                    if right_dd:
                        ps = right_dd.find_all("p")
                        if ps:
                            correct_answer = "\n".join(p.get_text(strip=True) for p in ps)
                        else:
                            correct_answer = right_dd.get_text(strip=True)
                        correct_answer = re.sub(r'<[^>]+>', '', correct_answer).strip()
                if not correct_answer:
                    b_daan = mark_answer.find("p", class_="B_daan")
                    if b_daan:
                        correct_answer = b_daan.get_text(strip=True)
            return Question(
                answer_type=answer_type,
                question_title=question_title,
                question_answers=options if options else None,
                correct_answer=correct_answer
            )
        except:
            return None

    def get_questions_from_html(self, html_content: str) -> List[Question]:
        try:
            soup = BeautifulSoup(html_content, "lxml")
            question_divs = soup.find_all("div", class_="questionLi")
            if not question_divs:
                question_divs = soup.find_all("div", class_="singleQuesId")
            questions = []
            for q_div in question_divs:
                question = self._parse_question_div(q_div)
                if question:
                    questions.append(question)
            return questions
        except:
            return []


class WordExporter:
    def __init__(self, course_name: str):
        self.course_name = course_name
        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)

    def export_word(self, questions: List[Question]):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
        except ImportError:
            print("python-docx库未安装，无法导出Word文档")
            return
        filename = self.output_dir / f"{self.course_name}_自测题库.docx"
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.add_heading(f"{self.course_name} 自测题库", level=0)
            type_names = {v: k for k, v in ExamCrawler.ANSWER_TYPES.items()}
            for i, q in enumerate(questions, 1):
                q_type = type_names.get(q.answer_type, "未知题型")
                para = doc.add_paragraph()
                run = para.add_run(f"{i}. {q.question_title}")
                run.bold = True
                type_para = doc.add_paragraph()
                type_run = type_para.add_run(f"【{q_type}】")
                type_run.font.size = Pt(10)
                if q.question_answers:
                    for opt in q.question_answers:
                        p = doc.add_paragraph(opt.strip(), style="List Bullet")
                if q.correct_answer:
                    ans_para = doc.add_paragraph()
                    ans_run = ans_para.add_run(f"答案: {q.correct_answer}")
                    ans_run.bold = True
                doc.add_paragraph()
            doc.save(filename)
            print(f"Word导出完成: {filename}")
        except Exception as e:
            print(f"Word导出失败: {e}")

    def export_markdown(self, questions: List[Question]):
        filename = self.output_dir / f"{self.course_name}_自测题库.md"
        try:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(f"# {self.course_name} 自测题库\n\n")
                type_names = {v: k for k, v in ExamCrawler.ANSWER_TYPES.items()}
                for i, q in enumerate(questions, 1):
                    q_type = type_names.get(q.answer_type, "未知题型")
                    f.write(f"## {i}. {q.question_title}\n\n")
                    f.write(f"**【{q_type}】**\n\n")
                    if q.question_answers:
                        for opt in q.question_answers:
                            f.write(f"- {opt.strip()}\n")
                        f.write("\n")
                    if q.correct_answer:
                        f.write(f"**答案**: {q.correct_answer}\n\n")
                    f.write("---\n\n")
            print(f"Markdown导出完成: {filename}")
        except Exception as e:
            print(f"Markdown导出失败: {e}")


def main():
    parser = argparse.ArgumentParser(description='学习通自测题目抓取工具')
    parser.add_argument('phone', nargs='?', help='学习通账号(手机号)')
    parser.add_argument('password', nargs='?', help='密码')
    parser.add_argument('--course', '-c', type=int, help='课程编号')
    parser.add_argument('--exam', '-e', type=int, help='自测编号')
    args = parser.parse_args()

    print("=" * 50)
    print("学习通自测题目抓取工具")
    print("=" * 50)

    crawler = ExamCrawler()
    course_name = ""

    if args.phone and args.password:
        if not crawler.login(args.phone, args.password):
            print("登录失败")
            return
        courses = crawler.get_courses()
        if not courses:
            print("未获取到课程列表")
            return
        print("\n课程列表:")
        for i, course in enumerate(courses, 1):
            print(f"({i}). {course.course_name}")
        if args.course:
            course_num = args.course - 1
        else:
            course_num = int(input("\n请选择课程编号: ").strip()) - 1
        if course_num < 0 or course_num >= len(courses):
            print("无效的课程编号")
            return
        selected_course = courses[course_num]
        print(f"\n选择课程: {selected_course.course_name}")
        exams = crawler.get_exams(selected_course)
        if not exams:
            print("未获取到自测列表")
            return
        print("\n自测列表:")
        for i, exam in enumerate(exams, 1):
            print(f"({i}). {exam.exam_name}")
        if args.exam:
            exam_num = args.exam - 1
        else:
            exam_num = int(input("\n请选择自测编号: ").strip()) - 1
        if exam_num < 0 or exam_num >= len(exams):
            print("无效的自测编号")
            return
        selected_exam = exams[exam_num]
        print(f"\n选择自测: {selected_exam.exam_name}")
        course_name = selected_exam.exam_name
        questions = crawler.get_exam_questions(selected_exam)
    else:
        print("\n登录方式:")
        print("(1). 账号密码登录")
        print("(2). 跳过登录，直接解析本地HTML文件")
        login_choice = input("\n请选择登录方式: ").strip()
        if login_choice == "1":
            phone = input("请输入学习通账号(手机号): ").strip()
            password = input("请输入密码: ").strip()
            if not phone or not password:
                print("账号和密码不能为空")
                return
            if not crawler.login(phone, password):
                print("登录失败")
                return
            courses = crawler.get_courses()
            if not courses:
                print("未获取到课程列表")
                return
            print("\n课程列表:")
            for i, course in enumerate(courses, 1):
                print(f"({i}). {course.course_name}")
            course_num = int(input("\n请选择课程编号: ").strip()) - 1
            if course_num < 0 or course_num >= len(courses):
                print("无效的课程编号")
                return
            selected_course = courses[course_num]
            print(f"\n选择课程: {selected_course.course_name}")
            exams = crawler.get_exams(selected_course)
            if not exams:
                print("未获取到自测列表")
                return
            print("\n自测列表:")
            for i, exam in enumerate(exams, 1):
                print(f"({i}). {exam.exam_name}")
            exam_num = int(input("\n请选择自测编号: ").strip()) - 1
            if exam_num < 0 or exam_num >= len(exams):
                print("无效的自测编号")
                return
            selected_exam = exams[exam_num]
            print(f"\n选择自测: {selected_exam.exam_name}")
            course_name = selected_exam.exam_name
            questions = crawler.get_exam_questions(selected_exam)
        elif login_choice == "2":
            html_file = input("请输入本地HTML文件路径: ").strip()
            course_name = input("请输入课程名称(用于文件名): ").strip()
            if not html_file:
                print("文件路径不能为空")
                return
            if not course_name:
                course_name = "自测题库"
            if not Path(html_file).exists():
                print(f"文件不存在: {html_file}")
                return
            with open(html_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            questions = crawler.get_questions_from_html(html_content)
        else:
            print("无效的选择")
            return

    if not questions:
        print("未获取到任何题目")
        return

    exporter = WordExporter(course_name)
    exporter.export_word(questions)
    exporter.export_markdown(questions)
    print("\n所有任务完成！")


if __name__ == "__main__":
    main()