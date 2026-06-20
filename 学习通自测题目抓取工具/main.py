"""
超星学习通作业爬取工具
支持登录、课程选择、作业爬取和导出功能
"""

import base64
import json
import logging
import re
import time
from argparse import ArgumentParser
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from urllib import parse
from urllib.parse import urlparse
try:
    import yaml
except ImportError:
    yaml = None

import requests
from bs4 import BeautifulSoup
from bs4.element import Tag
from Crypto.Cipher import AES
from Crypto.Util.Padding import pad, unpad

import os
from PIL import Image
from fpdf import FPDF
from concurrent.futures import ThreadPoolExecutor
import tempfile

# OpenAI兼容API相关导入
try:
    import openai
except ImportError:
    openai = None

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('fanya_crawler.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class Question:
    """题目数据类"""
    answer_type: int
    question_title: str
    correct_answer: Union[str, List[str]]
    question_answers: Optional[List[str]] = None
    ai_generated_answer: Optional[str] = None  # AI生成的答案
    ai_answer_confirmed: bool = False  # 是否已确认AI答案


@dataclass
class Assignment:
    """作业数据类"""
    work_id: str
    assignment_name: str
    assignment_status: str
    assignment_url: str
    course_id: str
    questions: List[Question] = field(default_factory=list)


@dataclass
class Course:
    """课程数据类"""
    course_id: str
    class_id: str
    cpi: str
    course_name: str
    course_url: str


@dataclass
class Resource:
    """资料数据类"""
    dataname: str
    datatype: str
    dataid: str


class FanyaCrawlerError(Exception):
    """自定义异常类"""
    pass


class AESCrypto:
    """AES加密解密工具类"""

    def __init__(self, key: str = "u2oh6Vu^HWe4_AES"):
        self.key = key.encode('utf-8')

    def encrypt(self, message: str) -> str:
        """AES加密"""
        try:
            cipher = AES.new(self.key, AES.MODE_CBC, self.key)
            padded_message = pad(message.encode('utf-8'),
                                 AES.block_size, style='pkcs7')
            encrypted_bytes = cipher.encrypt(padded_message)
            return base64.b64encode(encrypted_bytes).decode('utf-8')
        except Exception as e:
            logger.error(f"加密失败: {e}")
            raise FanyaCrawlerError(f"加密失败: {e}")

    def decrypt(self, encrypted_str: str) -> Tuple[bool, str]:
        """AES解密"""
        try:
            encrypted_bytes = base64.b64decode(encrypted_str)
            cipher = AES.new(self.key, AES.MODE_CBC, self.key)
            decrypted_bytes = unpad(cipher.decrypt(
                encrypted_bytes), AES.block_size)
            return True, decrypted_bytes.decode('utf-8')
        except Exception as e:
            logger.error(f"解密失败: {e}")
            return False, str(e)


class AIQuestionSolver:
    """AI题目解析器，支持OpenAI兼容的API"""

    def __init__(self, api_key: str, base_url: Optional[str] = None, model: str = ""):
        if openai is None:
            raise ImportError("openai库未安装，请运行 'pip install openai' 安装")

        self.api_key = api_key
        self.model = model

        # 配置OpenAI客户端
        if base_url:
            self.client = openai.OpenAI(api_key=api_key, base_url=base_url)
        else:
            self.client = openai.OpenAI(api_key=api_key)

    def solve_question(self, question: Question) -> str:
        """使用AI解决单个题目"""
        try:
            # 根据题目类型构建提示词
            prompt = self._build_prompt(question)

            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "你是一个专业的教育助手，擅长解答各种类型的题目。请根据题目内容给出准确、简洁的答案。同时请注意不要使用任何 Markdown 语法，只使用文本排版即可。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,  # 降低温度以获得更一致的答案
                max_tokens=500
            )

            ai_answer = response.choices[0].message.content.strip()
            logger.info(f"AI已解答题目: {question.question_title[:50]}...")
            return ai_answer

        except Exception as e:
            logger.error(f"AI解答题目失败: {e}")
            return f"AI解答失败: {str(e)}"

    def _build_prompt(self, question: Question) -> str:
        """根据题目类型构建AI提示词"""
        # 确定题目类型
        type_names = {v: k for k, v in FanyaCrawler.ANSWER_TYPES.items()}
        question_type = type_names.get(question.answer_type, "未知题型")

        prompt = f"题目类型: {question_type}\n"
        prompt += f"题目: {question.question_title}\n"

        # 如果是选择题，添加选项
        if question.question_answers:
            prompt += "选项:\n"
            for i, option in enumerate(question.question_answers, 1):
                prompt += f"{i}. {option.strip()}\n"

        prompt += "\n请给出这道题的正确答案，并简要说明解题思路。"

        return prompt


class ConfigManager:
    """配置管理器，处理配置文件的读取和保存"""

    def __init__(self, config_path: Optional[str] = None):
        if yaml is None:
            logger.warning("PyYAML库未安装，配置文件功能不可用。请运行 'pip install pyyaml' 安装。")
            self.config_path = None
            self.config = {}
            return

        self.config_path = Path(
            config_path) if config_path else Path("config.yaml")
        self.config = self.load_config()

    def load_config(self) -> Dict:
        """加载配置文件"""
        if self.config_path is None:
            return {}

        if self.config_path.exists():
            try:
                with open(self.config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f) or {}
                logger.info(f"配置文件已加载: {self.config_path}")
                return config
            except Exception as e:
                logger.error(f"加载配置文件失败: {e}")
                return {}
        else:
            logger.info("配置文件不存在，将使用默认配置")
            return {}

    def save_config(self, config: Dict):
        """保存配置文件"""
        if self.config_path is None:
            logger.warning("配置文件功能不可用，无法保存配置")
            return

        try:
            # 确保目录存在
            self.config_path.parent.mkdir(parents=True, exist_ok=True)

            with open(self.config_path, 'w', encoding='utf-8') as f:
                yaml.safe_dump(
                    config, f, default_flow_style=False, allow_unicode=True)
            logger.info(f"配置文件已保存: {self.config_path}")
        except Exception as e:
            logger.error(f"保存配置文件失败: {e}")

    def get_ai_config(self) -> Dict:
        """获取AI配置"""
        if yaml is None:
            return {
                'api_key': '',
                'base_url': '',
                'model': '',
                'enabled': False
            }

        ai_config = self.config.get('ai', {})
        return {
            'api_key': ai_config.get('api_key', ''),
            'base_url': ai_config.get('base_url', ''),
            'model': ai_config.get('model', ''),
            'enabled': ai_config.get('enabled', False)
        }

    def set_ai_config(self, api_key: str = '', base_url: str = '', model: str = '', enabled: bool = False):
        """设置AI配置"""
        if yaml is None:
            logger.error("PyYAML库未安装，无法保存配置。请运行 'pip install pyyaml' 安装。")
            return

        if 'ai' not in self.config:
            self.config['ai'] = {}

        self.config['ai'].update({
            'api_key': api_key,
            'base_url': base_url,
            'model': model,
            'enabled': enabled
        })

        self.save_config(self.config)


class FanyaCrawler:
    """超星学习通爬虫主类"""

    # 答题类型映射
    ANSWER_TYPES = {
        "单选题": 1,
        "多选题": 2,
        "填空题": 3,
        "判断题": 4,
        "思维导图": 5,
        "名词解释": 6,
        "简答题": 7,
        "其它": 255
    }

    # API endpoints
    API_ENDPOINTS = {
        'login': 'http://passport2.chaoxing.com/fanyalogin',
        'course_list': 'https://mooc2-ans.chaoxing.com/mooc2-ans/visit/courselistdata',
        'course_middle': 'https://mooc1.chaoxing.com/visit/stucoursemiddle',
        'work_list': 'https://mooc1.chaoxing.com/mooc2/work/list',
        'work_view': 'https://mooc1.chaoxing.com/mooc-ans/mooc2/work/view',
        'resources_list': 'https://mooc2-ans.chaoxing.com/mooc2-ans/coursedata/stu-datalist',
        'update_readcount': 'https://mooc2-ans.chaoxing.com/mooc2-ans/coursedata/update/update-read-count',
        'get_wps_preview': 'https://mooc2-ans.chaoxing.com/mooc2-ans/coursedata/get-preview-url',
    }

    SUPPORTED_FILE_TYPES = [
        "ppt",
        "pptx",
        "pdf"
    ]

    def __init__(self, ai_solver: Optional[AIQuestionSolver] = None):
        self.session = requests.Session()
        self.crypto = AESCrypto()
        self.ai_solver = ai_solver  # AI解析器
        self._setup_session()

    def _setup_session(self):
        """配置会话"""
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
            '(KHTML, like Gecko) Chrome/132.0.0.0 Safari/537.36 Edg/132.0.0.0'
        })

    def login(self, phone: str, password: str) -> bool:
        """登录功能"""
        try:
            encrypted_phone = self.crypto.encrypt(phone)
            encrypted_password = self.crypto.encrypt(password)

            headers = {
                "Accept": "application/json, text/javascript, */*; q=0.01",
                "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
                "Origin": "https://passport2.chaoxing.com",
                "Referer": "https://passport2.chaoxing.com/login",
                "X-Requested-With": "XMLHttpRequest"
            }

            payload = {
                'uname': encrypted_phone,
                'password': encrypted_password,
                't': "true",
                'validate': None,
                'doubleFactorLogin': "0",
                'independentId': "0"
            }

            response = self.session.post(
                self.API_ENDPOINTS['login'],
                data=payload,
                headers=headers,
                timeout=30
            )
            response.raise_for_status()

            result = response.json()
            if result.get('status'):
                logger.info("登录成功")
                return True
            else:
                logger.error(f"登录失败: {result.get('msg', '未知错误')}")
                return False

        except requests.RequestException as e:
            logger.error(f"登录请求失败: {e}")
            return False
        except Exception as e:
            logger.error(f"登录异常: {e}")
            return False

    def get_courses(self) -> List[Course]:
        """获取课程列表"""
        try:
            params = {
                "courseType": "1",
                "courseFolderId": "0",
                "superstarClass": "0"
            }

            response = self.session.get(
                self.API_ENDPOINTS['course_list'],
                params=params,
                timeout=30
            )
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "lxml")
            course_divs = soup.find_all("div", attrs={"class": "course-info"})

            courses = []
            for course_div in course_divs:
                try:
                    link = course_div.find_next("a", attrs={"class": "color1"})
                    if not link:
                        continue

                    course_url = link.get("href", "")
                    name_span = link.find_next(
                        "span", attrs={"class": "course-name overHidden2"})
                    if not name_span:
                        continue

                    course_name = name_span.get_text(strip=True)

                    # 解析URL参数
                    parsed_url = urlparse(course_url)
                    url_data = parse.parse_qs(parsed_url.query)

                    course = Course(
                        course_id=url_data.get("courseid", [""])[0],
                        class_id=url_data.get("clazzid", [""])[0],
                        cpi=url_data.get("cpi", [""])[0],
                        course_name=course_name,
                        course_url=course_url
                    )

                    if course.course_id and course.class_id:
                        courses.append(course)

                except Exception as e:
                    logger.warning(f"解析课程信息失败: {e}")
                    continue

            logger.info(f"获取到 {len(courses)} 门课程")
            return courses

        except Exception as e:
            logger.error(f"获取课程列表失败: {e}")
            raise FanyaCrawlerError(f"获取课程列表失败: {e}")

    def _get_workEnc(self, course: Course) -> str:
        try:
            # 获取课程中间页面
            middle_params = {
                "courseid": course.course_id,
                "clazzid": course.class_id,
                "cpi": course.cpi,
                "ismooc2": 1,
                "v": time.time(),
                "start": 0,
                "size": 500,
                "catalogId": 0,
                "superstarClass": 0,
            }

            middle_response = self.session.get(
                self.API_ENDPOINTS['course_middle'],
                params=middle_params,
                timeout=30,
                allow_redirects=True
            )
            middle_response.raise_for_status()

            # 提取作业编码参数
            soup = BeautifulSoup(middle_response.text, "lxml")
            work_enc_input = soup.find("input", id="workEnc")
            if not work_enc_input:
                logger.error("未找到工作加密参数")
                return ""

            return str(work_enc_input.get("value", ""))
        except Exception as e:
            logger.error(f"获取 workEnc 参数失败: {e}")
            raise FanyaCrawlerError(f"获取 workEnc 参数失败: {e}")

    def _get_EncAndT(self, course: Course) -> Tuple[str, str]:
        try:
            # 获取课程中间页面
            middle_params = {
                "courseid": course.course_id,
                "clazzid": course.class_id,
                "cpi": course.cpi,
                "ismooc2": 1,
                "v": time.time(),
                "start": 0,
                "size": 500,
                "catalogId": 0,
                "superstarClass": 0,
            }

            middle_response = self.session.get(
                self.API_ENDPOINTS['course_middle'],
                params=middle_params,
                timeout=30,
                allow_redirects=True
            )

            # 提取作业编码参数
            middle_response.raise_for_status()
            soup = BeautifulSoup(middle_response.text, "lxml")

            enc_tag = soup.find("input", {"id": "enc", "name": "enc"})
            enc = ""
            if isinstance(enc_tag, Tag):
                enc = enc_tag.get("value")
            else:
                raise FanyaCrawlerError(f"无法获取 enc 参数 input 标签: {enc_tag}")

            t_tag = soup.find("input", {"id": "t", "name": "t"})
            t = ""
            if isinstance(t_tag, Tag):
                t = t_tag.get("value")
            else:
                raise FanyaCrawlerError(f"无法获取 t 参数 input 标签: {t_tag}")

            return (enc, t)

        except Exception as e:
            logger.error(f"获取 enc 参数失败: {e}")
            raise FanyaCrawlerError(f"获取 enc 参数失败: {e}")

    def get_assignments(self, course: Course) -> List[Assignment]:
        """获取课程作业列表"""
        try:
            # 获取作业列表
            work_params = {
                "courseId": course.course_id,
                "classId": course.class_id,
                "cpi": course.cpi,
                "ut": "s",
                "enc": self._get_workEnc(course),
            }

            assignments = []
            page_num = 1
            total_page = 1

            while True:
                work_params["pageNum"] = page_num

                response = self.session.get(
                    self.API_ENDPOINTS['work_list'],
                    params=work_params,
                    timeout=30
                )
                response.raise_for_status()

                soup = BeautifulSoup(response.text, "lxml")

                if page_num == 1:
                    pagenum_pattern = r'pageNum\s*:\s*(\d+)'
                    scripts_in_soup = soup.find("body").find("script")
                    jscode_in_soup = scripts_in_soup.string if scripts_in_soup else ""
                    total_pagenum = re.search(pagenum_pattern, jscode_in_soup)

                    if (total_pagenum):
                        total_page = int(total_pagenum.group(1))

                li_tags = soup.find_all("li")

                if not li_tags:
                    break

                page_assignments = []
                for li in li_tags:
                    try:
                        data_url = li.get("data", "")
                        if not data_url:
                            continue

                        parsed_url = urlparse(data_url)
                        url_data = parse.parse_qs(parsed_url.query)

                        work_id = url_data.get("workId", [""])[0]
                        if not work_id:
                            continue

                        # 获取作业名称和状态
                        name_p = li.find("p")
                        assignment_name = name_p.get_text(
                            strip=True) if name_p else "未知作业"

                        status_p = name_p.find_next("p") if name_p else None
                        assignment_status = status_p.get_text(
                            strip=True) if status_p else "未知状态"

                        assignment = Assignment(
                            work_id=work_id,
                            assignment_name=assignment_name,
                            assignment_status=assignment_status,
                            assignment_url=data_url,
                            course_id=course.course_id
                        )

                        page_assignments.append(assignment)

                    except Exception as e:
                        logger.warning(f"解析作业信息失败: {e}")
                        continue

                assignments.extend(page_assignments)

                # 检查是否还有下一页
                if (total_page == page_num):
                    break

                if (total_page != 1):
                    page_num += 1

            logger.info(f"获取到 {len(assignments)} 个作业")
            return assignments

        except Exception as e:
            logger.error(f"获取作业列表失败: {e}")
            raise FanyaCrawlerError(f"获取作业列表失败: {e}")

    def get_resource_list(self, course: Course) -> Dict[str, Dict[int, str]]:
        ret = dict()

        self.current_enc, self.current_time = self._get_EncAndT(course)

        def get_folder_detail(params: Dict[str, str]):
            try:
                response = self.session.get(self.API_ENDPOINTS['resources_list'],
                                            params=stu_datalist_payload,
                                            timeout=30)
                response.raise_for_status()

                # logger.info(f"Current request header: {response.request.headers}")

                soup = BeautifulSoup(response.text, "lxml")
                root_dataBody = soup.find("div", {"class": "dataBody"})
                if isinstance(root_dataBody, Tag):
                    root_dir = list()
                    root_directory = root_dataBody.find_all(
                        "ul", {"class": "dataBody_td"})
                    for dir in root_directory:
                        root_dir.append(Resource(
                                        dataname=dir.get("dataname"),
                                        datatype=dir.get("type"),
                                        dataid=dir.get("id")
                                        ))
                    return root_dir
                return list()
            except Exception as e:
                logger.error(f"获取文件列表失败: {e}")
                return list()

        def parse_resourse(resource: Resource):
            logger.info(f"正在准备解析资源: {resource.dataname}")
            # 初始化请求参数
            readcount_payload = {
                "courseId": course.course_id,
                "clazzId": course.class_id,
                "dataId": resource.dataid,
                "cpi": course.cpi,
                "ut": "s"
            }

            # 发起阅读量计数请求
            response = self.session.get(
                self.API_ENDPOINTS['update_readcount'],
                params=readcount_payload,
                timeout=30,
                allow_redirects=True
            )

            update_status = json.loads(response.text)
            if update_status["msg"] == "success":
                logger.info("阅读量计数回调成功，开始获取文件")
            else:
                logger.error(f"阅读量计数失败: {response.text}")
                raise FanyaCrawlerError(f"阅读量计数失败: {response.text}")

            preview_payload = {
                "courseid": course.course_id,
                "clazzid": course.class_id,
                "dataId": resource.dataid,
                "cpi": course.cpi,
                "ut": "s"
            }

            # 发起文件获取请求
            response = self.session.get(
                self.API_ENDPOINTS['get_wps_preview'],
                params=preview_payload,
                timeout=30,
                allow_redirects=True
            )

            file_status = json.loads(response.text)
            preview_url = ""
            if file_status["status"]:
                preview_url = file_status["url"]
                logger.info(f"成功获取文件预览 URL: {preview_url}，准备抓取")
                image_urls = dict()
                response = self.session.get(
                    preview_url,
                    timeout=30,
                    allow_redirects=True)

                soup = BeautifulSoup(response.text, "lxml")
                filebox = soup.find("div", {"class": "fileBox"})
                if isinstance(filebox, Tag):
                    all_image_tag = filebox.find_all("li")
                    for li in all_image_tag:
                        page = int(li.find("span").text)
                        imgsrc = li.find("img").get("src")
                        image_urls[page] = imgsrc
                    return image_urls

            else:
                logger.error(f"获取文件预览失败: {response.text}")
                logger.error(f"Request URL: {response.request.url}")
                raise FanyaCrawlerError(f"获取文件预览失败: {response.text}")

        # 构造请求头
        stu_datalist_payload = {
            "courseid": course.course_id,
            "clazzid": course.class_id,
            "cpi": course.cpi,
            "ut": "s",
            "t": self.current_time,
            "stuenc": self.current_enc,
        }

        # 尝试获取文件列表
        while True:
            try:
                root_dir = get_folder_detail(stu_datalist_payload)
                if len(root_dir):
                    logger.info("获取根目录成功")
                    print("\n请选择需要下载的文件或需要访问的目录:")
                    print("\n如果有多个文件需要下载，请使用英文逗号分隔，如果是连续的多个文件，请使用 \'-\' 连接")

                    for i in range(1, len(root_dir) + 1):
                        print(f"({i}). {root_dir[i - 1].dataname}")

                    user_choice = str(
                        input(f"\n请输入 1 至 {len(root_dir)} 之间的数字: "))
                    while True:
                        if '-' in user_choice:
                            interval = user_choice.split("-")
                            current_todo = [x for x in range(
                                int(interval[0]), int(interval[1]) + 1)]
                            if len(current_todo) > len(root_dir):
                                raise FanyaCrawlerError(
                                    f"选择数目 {len(current_todo)} 大于目录项数 {len(root_dir)}")
                            break
                        elif ',' in user_choice:
                            current_todo = [int(x)
                                            for x in user_choice.split(',')]
                            if max(current_todo) > len(root_dir):
                                raise FanyaCrawlerError(f"选择不在范围内")
                            break
                        else:
                            current_todo = int(user_choice)
                            if current_todo > len(root_dir):
                                raise FanyaCrawlerError(f"选择不在范围内")
                            break

                    print(f"\n你的选择是: {current_todo}")
                    if isinstance(current_todo, int):
                        if root_dir[current_todo - 1].datatype == "afolder":
                            logger.info("检测到用户选择了文件夹，正在切换目录")
                            stu_datalist_payload = {
                                "courseid": course.course_id,
                                "dataName": root_dir[current_todo - 1].dataname,
                                "dataId": root_dir[current_todo - 1].dataid,
                                "type": 1,
                                "flag": 0,
                                "clazzid": course.class_id,
                                "enc": self.current_enc,
                                "ut": "s",
                                "t": self.current_time,
                                "cpi": course.cpi,
                                "microTopicId": 0
                            }
                            continue
                        elif root_dir[current_todo - 1].datatype in self.SUPPORTED_FILE_TYPES:
                            ret[root_dir[current_todo -
                                1].dataname] = parse_resourse(root_dir[current_todo - 1])
                            break
                    elif isinstance(current_todo, list):
                        for index in current_todo:
                            if root_dir[index - 1].datatype == "afolder":
                                logger.error("不支持多选文件夹")
                                raise FanyaCrawlerError("不支持多选文件夹")
                            if root_dir[index - 1].datatype in self.SUPPORTED_FILE_TYPES:
                                ret[root_dir[index -
                                    1].dataname] = parse_resourse(root_dir[index - 1])
                        break
                else:
                    logger.error("获取根目录失败，可能该课程没有上传资料！")
                    return dict()

            except Exception as e:
                logger.error(f"获取文件列表页面失败: {e}")
                raise FanyaCrawlerError(f"获取文件列表失败: {e}")
        logger.info(f"成功获取了 {len(ret)} 个资源")
        return ret

    def _normalize_title(self, title: str) -> str:
        """标准化题目标题"""
        return (title.strip()
                .replace("（", "(")
                .replace("）", ")")
                .replace("\u200c", "")
                .replace("\u200e", "")
                .replace("\u200d", "")
                .replace("\u200f", "")
                .replace("\xa0", ""))

    def _normalize_answers(self, answers_text: str) -> List[str]:
        """标准化答案选项"""
        answer_list = answers_text.split("\n")
        return [answer.strip() + "\n" for answer in answer_list if answer.strip()]

    def _parse_questions(self, question_block: Tag) -> List[Question]:
        """解析题目块"""
        try:
            title_tag = question_block.find("h2", attrs={"class": "type_tit"})
            if not title_tag:
                return []

            block_title = title_tag.get_text(strip=True)

            # 确定题目类型
            answer_type = 0
            for type_name, type_id in self.ANSWER_TYPES.items():
                if type_name in block_title:
                    answer_type = type_id
                    break

            if answer_type == 0:
                logger.warning(f"未知题目类型: {block_title}")
                return []

            logger.info(f"解析题目类型: {block_title}")

            # 获取所有题目详情
            question_details = question_block.find_all(
                "div", attrs={"aria-label": "题目详情"})
            questions = []

            for detail in question_details:
                try:
                    question = self._parse_single_question(detail, answer_type)
                    if question:
                        questions.append(question)
                except Exception as e:
                    logger.warning(f"解析单个题目失败: {e}")
                    continue

            return questions

        except Exception as e:
            logger.error(f"解析题目块失败: {e}")
            return []

    def _parse_single_question(self, detail: Tag, answer_type: int) -> Optional[Question]:
        """解析单个题目"""
        try:
            # 获取题目标题
            title_tag = detail.find(
                "h3", attrs={"class": "mark_name colorDeep"})
            if not title_tag:
                return None

            question_title = self._normalize_title(title_tag.get_text())

            if answer_type in [self.ANSWER_TYPES["单选题"], self.ANSWER_TYPES["多选题"]]:
                # 选择题
                answers_tag = detail.find(
                    "ul", attrs={"class": "mark_letter colorDeep qtDetail"})
                question_answers = self._normalize_answers(
                    answers_tag.get_text()) if answers_tag else []

                answer_div = detail.find("div", attrs={"class": "mark_answer"})
                correct_tag = answer_div.find(
                    "span", attrs={"class": "rightAnswerContent workTextWrap"}) if answer_div else None
                correct_answer = correct_tag.get_text(
                    strip=True) if correct_tag else ""

                return Question(
                    answer_type=answer_type,
                    question_title=question_title,
                    question_answers=question_answers,
                    correct_answer=correct_answer
                )

            elif answer_type == self.ANSWER_TYPES["填空题"]:
                # 填空题
                fill_tag = detail.find(
                    "dl", attrs={"class": "mark_fill colorGreen"})
                if fill_tag:
                    dd_tags = fill_tag.find_all("dd")
                    correct_answers = [dd.get_text(
                        strip=True) for dd in dd_tags]
                else:
                    correct_answers = []

                return Question(
                    answer_type=answer_type,
                    question_title=question_title,
                    correct_answer=correct_answers
                )

            elif answer_type == self.ANSWER_TYPES["判断题"]:
                # 判断题
                answer_div = detail.find("div", attrs={"class": "mark_answer"})
                correct_tag = answer_div.find(
                    "span", attrs={"class": "rightAnswerContent"}) if answer_div else None
                correct_answer = correct_tag.get_text(
                    strip=True) if correct_tag else ""

                return Question(
                    answer_type=answer_type,
                    question_title=question_title,
                    correct_answer=correct_answer
                )

            elif answer_type in [self.ANSWER_TYPES["名词解释"], self.ANSWER_TYPES["简答题"]]:
                # 名词解释和简答题
                answer_div = detail.find("div", attrs={"class": "mark_answer"})
                if answer_div:
                    # 尝试获取正确答案，如果不可用则获取学生答案
                    correct_tag = answer_div.find(
                        "span", attrs={"class": "rightAnswerContent"})
                    if correct_tag:
                        # 如果有正确答案
                        correct_answer = correct_tag.get_text(strip=True)
                    else:
                        # 获取学生答案作为替代
                        stu_answer_tag = answer_div.find(
                            "dd", attrs={"class": "textwrap stuAnswerContent reserve-newline"}
                        )
                        if stu_answer_tag:
                            correct_answer = stu_answer_tag.get_text(
                                strip=True)
                        else:
                            correct_answer = ""
                else:
                    correct_answer = ""

                return Question(
                    answer_type=answer_type,
                    question_title=question_title,
                    correct_answer=correct_answer
                )

        except Exception as e:
            logger.warning(f"解析单个题目详情失败: {e}")

        return None

    def get_assignment_questions(self, assignment: Assignment) -> List[Question]:
        """获取作业题目"""
        try:
            parsed_url = urlparse(assignment.assignment_url)
            url_data = parse.parse_qs(parsed_url.query)

            params = {
                "courseId": url_data.get("courseId", [""])[0],
                "classId": url_data.get("classId", [""])[0],
                "cpi": url_data.get("cpi", [""])[0],
                "workId": url_data.get("workId", [""])[0],
                "answerId": url_data.get("answerId", [""])[0],
                "enc": url_data.get("enc", [""])[0]
            }

            response = self.session.get(
                self.API_ENDPOINTS['work_view'],
                params=params,
                timeout=30
            )
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "lxml")
            question_blocks = soup.find_all(
                "div", attrs={"class": "mark_item"})

            all_questions = []
            for block in question_blocks:
                questions = self._parse_questions(block)
                all_questions.extend(questions)

            logger.info(
                f"作业 {assignment.assignment_name} 获取到 {len(all_questions)} 道题目")
            return all_questions

        except Exception as e:
            logger.error(f"获取作业题目失败: {e}")
            return []

    def solve_assignment_questions_with_ai(self, assignment: Assignment) -> List[Question]:
        """使用AI解决作业中的题目"""
        if not self.ai_solver:
            logger.warning("AI解析器未初始化，跳过AI解答")
            return assignment.questions

        logger.info(f"开始使用AI解答作业: {assignment.assignment_name}")

        for i, question in enumerate(assignment.questions):
            logger.info(
                f"正在使用AI解答第 {i+1}/{len(assignment.questions)} 题: {question.question_title[:30]}...")

            # 使用AI解答题目
            ai_answer = self.ai_solver.solve_question(question)
            question.ai_generated_answer = ai_answer
            question.ai_answer_confirmed = True  # 标记AI答案已生成

            time.sleep(0.5)  # 避免请求过于频繁

        logger.info(f"AI解答完成，共解答 {len(assignment.questions)} 道题目")
        return assignment.questions


class DocumentExporter:
    """文档导出器"""

    def __init__(self, course_name: str):
        self.course_name = course_name
        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)

    def export_markdown(self, assignments: List[Assignment], with_answers: bool = True, include_ai: bool = False):
        """导出Markdown格式"""
        suffix = "带答案"
        if include_ai:
            suffix += "_含AI解析"
        elif not with_answers:
            suffix = "不带答案"
        filename = self.output_dir / f"{self.course_name}_习题_{suffix}.md"

        try:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(f"# {self.course_name} 习题集\n\n")

                for assignment in assignments:
                    if not assignment.questions:
                        continue

                    f.write(f"## {assignment.assignment_name}\n\n")

                    for question in assignment.questions:
                        f.write(f"### {question.question_title}\n\n")

                        # 选择题显示选项
                        if question.question_answers:
                            for answer in question.question_answers:
                                f.write(answer)
                            f.write("\n")

                        # 显示答案或留空
                        if with_answers:
                            if isinstance(question.correct_answer, list):
                                f.write(
                                    "正确答案: " + ", ".join(question.correct_answer) + "\n\n")
                            else:
                                f.write(f"正确答案: {question.correct_answer}\n\n")

                            # 如果包含AI答案，也显示
                            if include_ai and question.ai_generated_answer:
                                f.write(
                                    f"AI解析: {question.ai_generated_answer}\n\n")
                        else:
                            f.write("答案: ____________________\n\n")

            logger.info(f"Markdown导出完成: {filename}")

        except Exception as e:
            logger.error(f"Markdown导出失败: {e}")

    def export_word(self, assignments: List[Assignment], with_answers: bool = True, include_ai: bool = False):
        """导出Word格式"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("python-docx库未安装，无法导出Word文档")
            return

        suffix = "带答案"
        if include_ai:
            suffix += "_含AI解析"
        elif not with_answers:
            suffix = "不带答案"
        filename = self.output_dir / f"{self.course_name}_习题_{suffix}.docx"

        try:
            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 添加标题
            doc.add_heading(f"{self.course_name} 习题集", level=0)

            for assignment in assignments:
                if not assignment.questions:
                    continue

                # 作业标题
                doc.add_heading(assignment.assignment_name, level=1)

                for i, question in enumerate(assignment.questions, 1):
                    # 题目标题
                    para = doc.add_paragraph()
                    run = para.add_run(f"{i}. {question.question_title}")
                    run.bold = True

                    # 选择题选项
                    if question.question_answers:
                        for answer in question.question_answers:
                            p = doc.add_paragraph(
                                answer.strip(), style="List Bullet")
                            p.paragraph_format.left_indent = Inches(0.5)

                    # 答案部分
                    if with_answers:
                        ans_para = doc.add_paragraph()
                        if isinstance(question.correct_answer, list):
                            ans_text = "正确答案: " + \
                                ", ".join(question.correct_answer)
                        else:
                            ans_text = f"正确答案: {question.correct_answer}"
                        run = ans_para.add_run(ans_text)
                        run.bold = True

                        # 如果包含AI答案，也显示
                        if include_ai and question.ai_generated_answer:
                            ai_para = doc.add_paragraph()
                            ai_run = ai_para.add_run(
                                f"AI解析: {question.ai_generated_answer}")
                            ai_run.italic = True
                    else:
                        doc.add_paragraph("答案: ____________________")

                    # 添加空行
                    doc.add_paragraph()

            doc.save(filename)
            logger.info(f"Word导出完成: {filename}")

        except Exception as e:
            logger.error(f"Word导出失败: {e}")

    def export_json(self, assignments: List[Assignment], include_ai: bool = False):
        """导出JSON格式（用于备份和数据交换）"""
        suffix = "_数据"
        if include_ai:
            suffix += "_含AI解析"
        filename = self.output_dir / f"{self.course_name}_习题{suffix}.json"

        try:
            data = {
                "course_name": self.course_name,
                "export_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                "include_ai": include_ai,
                "assignments": []
            }

            for assignment in assignments:
                assignment_data = {
                    "work_id": assignment.work_id,
                    "name": assignment.assignment_name,
                    "status": assignment.assignment_status,
                    "questions": []
                }

                for question in assignment.questions:
                    question_data = {
                        "type": question.answer_type,
                        "title": question.question_title,
                        "correct_answer": question.correct_answer,
                        "options": question.question_answers
                    }

                    # 如果包含AI答案，也添加到数据中
                    if include_ai:
                        question_data["ai_generated_answer"] = question.ai_generated_answer
                        question_data["ai_answer_confirmed"] = question.ai_answer_confirmed

                    assignment_data["questions"].append(question_data)

                data["assignments"].append(assignment_data)

            with open(filename, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            logger.info(f"JSON导出完成: {filename}")

        except Exception as e:
            logger.error(f"JSON导出失败: {e}")

    def exprot_pdf_from_url(self, image_urls: Dict[int, str], filename: str, max_workers=4):
        """
        将网络图片按照顺序组合成 PDF 文件

        参数:
        image_urls: dict - 键为页数，值为图片 URL 的字典
        output_filename: str - 输出 PDF 文件名
        max_workers: int - 并发下载图片的最大线程数
        """
        logger.info(f"即将开始下载 {filename} 的内容并生成 PDF")
        logger.info(f"将使用最大 {max_workers} workers 并行下载")
        # 按页数排序 URL
        sorted_urls = [url for _, url in sorted(
            image_urls.items(), key=lambda x: x[0])]

        # 创建临时目录存储下载的图片
        with tempfile.TemporaryDirectory() as temp_dir:
            # 下载所有图片（并发下载提高速度）
            def download_image(url_idx):
                idx, url = url_idx
                try:
                    response = requests.get(url, stream=True, timeout=30)
                    response.raise_for_status()

                    # 保存到临时文件
                    img_path = os.path.join(temp_dir, f"page_{idx}.jpg")
                    with open(img_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    return img_path
                except Exception as e:
                    logger.error(f"下载图片失败 (页数 {idx}, URL: {url}): {e}")
                    return None

            # 使用线程池并发下载
            image_paths = []
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                results = executor.map(
                    download_image, enumerate(sorted_urls, start=1))
                for result in results:
                    if result:
                        image_paths.append(result)

            # 按文件名中的页数排序图片路径
            image_paths.sort(key=lambda x: int(
                os.path.basename(x).split('_')[1].split('.')[0]))

            # 创建 PDF
            pdf = FPDF()

            for img_path in image_paths:
                try:
                    # 使用PIL检测图片尺寸
                    with Image.open(img_path) as img:
                        width_px, height_px = img.size

                        dpi = img.info.get('dpi', (72, 72))
                        if dpi[0] == 0 or dpi[1] == 0:  # 防止除零错误
                            dpi = (72, 72)

                        # 计算实际物理尺寸（单位：点）
                        width_pt = (width_px * 72) / dpi[0]
                        height_pt = (height_px * 72) / dpi[1]

                    # 添加新页面
                    pdf.add_page(format=(width_pt, height_pt))

                    # 添加图片到PDF
                    pdf.image(img_path, x=0, y=0, w=width_pt, h=height_pt)
                except Exception as e:
                    logger.error(f"处理图片失败 {img_path}: {e}")

            # 保存PDF
            pdf.output(os.path.join(self.output_dir, filename + ".pdf"))
            logger.info(
                f"PDF 已成功保存至: {os.path.join(self.output_dir, filename + ".pdf")}")


CRAWLER_OPERATIONS = {
    0: "爬取课后题",
    1: "爬取不提供下载的资料"
}


def main():
    """主函数"""
    parser = ArgumentParser(
        prog="fanya_crawler",
        description="超星学习通作业爬取工具"
    )
    parser.add_argument("phone", help="手机号")
    parser.add_argument("password", help="密码")
    parser.add_argument("--format", choices=["markdown", "word", "json", "all"],
                        default="all", help="导出格式")
    parser.add_argument("--no-answers", action="store_true", help="不包含答案")
    parser.add_argument("--config", help="配置文件路径", default="config.yaml")
    parser.add_argument("--setup-ai", action="store_true", help="设置AI配置")

    args = parser.parse_args()

    try:
        # 初始化配置管理器
        config_manager = ConfigManager(args.config)

        # 如果用户选择设置AI配置
        if args.setup_ai:
            if yaml is None:
                logger.error(
                    "PyYAML库未安装，无法使用配置文件功能。请运行 'pip install pyyaml' 安装。")
                return

            print("AI配置设置")
            print("=" * 30)

            current_config = config_manager.get_ai_config()

            # 获取用户输入
            api_key = input(
                f"请输入API密钥 (当前: {'*' * 20 if current_config.get('api_key') else '未设置'}): ").strip()
            if not api_key:
                api_key = current_config.get('api_key', '')

            base_url = input(
                f"请输入API基础URL (可选，当前: {current_config.get('base_url', '未设置')}): ").strip()
            if not base_url:
                base_url = current_config.get('base_url', '')

            model = input(f"请输入模型名称): ").strip()
            if not model:
                model = current_config.get('model', '')

            enabled = input(
                f"是否启用AI功能? (y/N, 当前: {current_config.get('enabled', 'False')}): ").strip().lower()
            if enabled == '':
                enabled = current_config.get('enabled', False)
            else:
                enabled = enabled in ['y', 'yes', 'true', '1']

            # 保存配置
            config_manager.set_ai_config(
                api_key=api_key,
                base_url=base_url,
                model=model,
                enabled=enabled
            )

            print("\nAI配置已保存！")
            return

        # 从配置文件获取AI配置
        ai_config = config_manager.get_ai_config()

        # 初始化AI解析器（如果配置了API密钥且启用了AI功能）
        ai_solver = None
        if ai_config.get('api_key') and ai_config.get('enabled', False):
            try:
                ai_solver = AIQuestionSolver(
                    api_key=ai_config['api_key'],
                    base_url=ai_config.get('base_url') or None,
                    model=ai_config.get('model', '')
                )
                logger.info("AI解析器初始化成功")
            except ImportError as e:
                logger.error(f"AI功能初始化失败: {e}")
                # 如果AI功能未正确安装，但用户期望使用AI，则退出
                if ai_config.get('enabled', False):
                    logger.error("由于AI功能未正确安装，程序将退出")
                    return
            except Exception as e:
                logger.error(f"AI解析器初始化失败: {e}")
                # 如果AI配置有问题，但用户期望使用AI，则退出
                if ai_config.get('enabled', False):
                    logger.error("程序将退出")
                    return

        # 初始化爬虫
        crawler = FanyaCrawler(ai_solver=ai_solver)

        # 登录
        logger.info("开始登录...")
        if not crawler.login(args.phone, args.password):
            logger.error("登录失败，程序退出")
            return

        # 获取课程列表
        logger.info("获取课程列表...")
        courses = crawler.get_courses()
        if not courses:
            logger.error("未找到课程，程序退出")
            return

        # 显示课程列表供用户选择
        print("\n请选择要爬取的课程:")
        for i, course in enumerate(courses, 1):
            print(f"{i}. {course.course_name}")

        while True:
            try:
                choice = int(input(f"\n请输入课程编号 (1-{len(courses)}): "))
                if 1 <= choice <= len(courses):
                    selected_course = courses[choice - 1]
                    break
                else:
                    print(f"请输入 1 到 {len(courses)} 之间的数字")
            except ValueError:
                print("请输入有效的数字")

        logger.info(f"选择课程: {selected_course.course_name}")

        while True:
            try:
                for key, opt in CRAWLER_OPERATIONS.items():
                    print(f'({key}).{opt}')
                current_operation = int(input(f"\n请选择您要爬取的数据: "))
                if 0 <= current_operation <= len(CRAWLER_OPERATIONS) - 1:
                    operation = current_operation
                    break
                else:
                    print(f"请输入 0 到 {len(CRAWLER_OPERATIONS) - 1} 之间的数字")
            except ValueError:
                print("请输入有效的数字")

        logger.info(f'Operation: {CRAWLER_OPERATIONS[operation]}')

        if operation == 0:
            # 获取作业列表
            logger.info("获取作业列表...")
            assignments = crawler.get_assignments(selected_course)
            if not assignments:
                logger.warning("该课程暂无作业")
                return

            # 获取每个作业的题目
            logger.info("开始爬取作业题目...")
            for assignment in assignments:
                logger.info(f"正在处理作业: {assignment.assignment_name}")
                questions = crawler.get_assignment_questions(assignment)
                assignment.questions = questions
                time.sleep(1)  # 避免请求过于频繁

            # 如果启用了AI功能，使用AI解答题目
            if ai_config.get('enabled', False) and crawler.ai_solver:
                logger.info("开始使用AI解答题目...")
                for assignment in assignments:
                    logger.info(f"正在AI解答作业: {assignment.assignment_name}")
                    crawler.solve_assignment_questions_with_ai(assignment)
                    time.sleep(1)  # 避免请求过于频繁

            # 导出文档
            exporter = DocumentExporter(selected_course.course_name)

            # 确定是否包含AI答案
            include_ai = ai_config.get(
                'enabled', False) and crawler.ai_solver is not None

            if args.format in ["markdown", "all"]:
                exporter.export_markdown(
                    assignments, with_answers=not args.no_answers, include_ai=include_ai)
                if not args.no_answers and not include_ai:
                    exporter.export_markdown(assignments, with_answers=False)

            if args.format in ["word", "all"]:
                exporter.export_word(
                    assignments, with_answers=not args.no_answers, include_ai=include_ai)
                if not args.no_answers and not include_ai:
                    exporter.export_word(
                        assignments, with_answers=False, include_ai=include_ai)

            if args.format in ["json", "all"]:
                exporter.export_json(assignments, include_ai=include_ai)

            logger.info("所有任务完成！")

        if operation == 1:
            logger.info("获取文件列表...")
            files = crawler.get_resource_list(selected_course)
            logger.info(f"成功获取文件: {list(files.keys())}")

            exporter = DocumentExporter(selected_course.course_name)

            for filename, file in files.items():
                exporter.exprot_pdf_from_url(file, filename)

    except KeyboardInterrupt:
        logger.info("用户中断程序")
    except Exception as e:
        logger.error(f"程序执行出错: {e}")
        raise


if __name__ == "__main__":
    main()